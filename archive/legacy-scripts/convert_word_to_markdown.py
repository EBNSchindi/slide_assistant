#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word-zu-Markdown-Konverter
Konvertiert .docx-Dateien exakt wortgetreu in Markdown-Format
"""

import argparse
import os
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE

# PERFORMANCE: Compile regex patterns once at module level instead of on every call
_HEADING_PATTERN = re.compile(r'heading\s*(\d+)', re.IGNORECASE)
_MULTIPLE_NEWLINES_PATTERN = re.compile(r'\n{3,}')


def format_runs(runs):
    """
    Formatiert Text-Runs mit Fett-, Kursiv-Formatierung und Links zu Markdown
    """
    result = []
    for run in runs:
        text = run.text
        if not text:
            continue
        
        # Hyperlink prüfen
        is_link = False
        link_url = None
        try:
            # Prüfe ob der Run Teil eines Hyperlinks ist
            parent = run._element.getparent()
            if parent is not None and parent.tag.endswith('hyperlink'):
                # Hyperlink-Element gefunden
                rel_id = parent.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                if rel_id:
                    # URL aus den Relationships extrahieren
                    try:
                        part = run.part
                        if hasattr(part, 'rels') and rel_id in part.rels:
                            rel = part.rels[rel_id]
                            link_url = rel.target_ref
                            is_link = True
                    except Exception:
                        # Alternative: Prüfe ob anchor-Attribut vorhanden ist
                        anchor = parent.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}anchor')
                        if anchor:
                            link_url = anchor
                            is_link = True
        except Exception:
            pass
        
        # Formatierungen anwenden
        formatted_text = text
        if run.bold:
            formatted_text = f"**{formatted_text}**"
        if run.italic:
            formatted_text = f"*{formatted_text}*"
        if run.underline:
            formatted_text = f"<u>{formatted_text}</u>"
        
        # Link-Formatierung
        if is_link and link_url:
            formatted_text = f"[{formatted_text}]({link_url})"
        
        result.append(formatted_text)
    
    return ''.join(result)


def get_heading_level(style_name):
    """
    Ermittelt die Überschriftenebene basierend auf dem Style-Namen
    """
    if not style_name:
        return None

    # PERFORMANCE: Use pre-compiled regex pattern
    if 'heading' in style_name.lower():
        match = _HEADING_PATTERN.search(style_name)
        if match:
            return int(match.group(1))

    return None


def process_paragraph(para):
    """
    Verarbeitet einen Word-Absatz und konvertiert ihn zu Markdown
    """
    style_name = para.style.name if para.style else None
    heading_level = get_heading_level(style_name)
    
    # Überschrift
    if heading_level:
        text = format_runs(para.runs)
        if text.strip():
            return f"{'#' * heading_level} {text}\n\n"
        return ""
    
    # Liste erkennen
    if para.style.name.startswith('List'):
        # Bullet-Liste
        if 'bullet' in para.style.name.lower() or para.style.name.startswith('List Bullet'):
            text = format_runs(para.runs)
            if text.strip():
                return f"- {text}\n"
        # Nummerierte Liste
        elif 'number' in para.style.name.lower() or para.style.name.startswith('List Number'):
            text = format_runs(para.runs)
            if text.strip():
                return f"1. {text}\n"
    
    # Prüfe auf Listen-Formatierung durch Einrückung
    if para.paragraph_format.left_indent and para.paragraph_format.left_indent.pt > 0:
        # Eingerückt - könnte Liste sein
        text = format_runs(para.runs)
        if text.strip():
            return f"  - {text}\n"
    
    # Normaler Absatz
    text = format_runs(para.runs)
    if text.strip():
        return f"{text}\n\n"
    
    return "\n"


def process_table(table):
    """
    Konvertiert eine Word-Tabelle zu Markdown-Tabelle
    """
    markdown_lines = []
    
    # Header-Zeile
    if table.rows:
        header_row = table.rows[0]
        header_cells = []
        for cell in header_row.cells:
            cell_text = ''.join(format_runs(cell.paragraphs[0].runs) if cell.paragraphs else '')
            header_cells.append(cell_text.strip())
        
        markdown_lines.append("| " + " | ".join(header_cells) + " |")
        markdown_lines.append("| " + " | ".join(["---"] * len(header_cells)) + " |")
        
        # Daten-Zeilen
        for row in table.rows[1:]:
            row_cells = []
            for cell in row.cells:
                cell_text = ''
                for para in cell.paragraphs:
                    cell_text += format_runs(para.runs)
                row_cells.append(cell_text.strip())
            markdown_lines.append("| " + " | ".join(row_cells) + " |")
    
    return "\n".join(markdown_lines) + "\n\n"


def convert_docx_to_markdown(docx_path, output_path=None):
    """
    Konvertiert eine .docx-Datei zu Markdown
    """
    try:
        doc = Document(docx_path)
    except Exception as e:
        print(f"Fehler beim Öffnen der Datei {docx_path}: {e}")
        return False

    markdown_content = []

    # PERFORMANCE: Pre-build element mappings to avoid O(n²) complexity
    # Old approach: For each element, iterate through all paragraphs/tables (O(n²))
    # New approach: Create hash map once (O(n)), then lookup in O(1)
    para_map = {p._element: p for p in doc.paragraphs}
    table_map = {t._element: t for t in doc.tables}

    # Durchlaufe alle Elemente im Dokument in der richtigen Reihenfolge
    for element in doc.element.body:
        # Absatz - O(1) lookup statt O(n) search
        if element.tag.endswith('p') and element in para_map:
            markdown_content.append(process_paragraph(para_map[element]))

        # Tabelle - O(1) lookup statt O(n) search
        elif element.tag.endswith('tbl') and element in table_map:
            markdown_content.append(process_table(table_map[element]))
    
    # Markdown-Text zusammenfügen
    markdown_text = ''.join(markdown_content)

    # PERFORMANCE: Use pre-compiled regex pattern
    # Überflüssige Leerzeilen reduzieren (max. 2 aufeinanderfolgende)
    markdown_text = _MULTIPLE_NEWLINES_PATTERN.sub('\n\n', markdown_text)
    
    # Whitespace am Ende entfernen
    markdown_text = markdown_text.rstrip() + '\n'
    
    # Ausgabe-Datei bestimmen
    if output_path is None:
        docx_file = Path(docx_path)
        output_path = docx_file.with_suffix('.md')
    
    # In Datei schreiben
    try:
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(markdown_text)
        print(f"Erfolgreich konvertiert: {docx_path} -> {output_path}")
        return True
    except Exception as e:
        print(f"Fehler beim Schreiben der Datei {output_path}: {e}")
        return False


def main():
    """
    Hauptfunktion für CLI-Interface
    """
    parser = argparse.ArgumentParser(
        description='Konvertiert Word-Dokumente (.docx) zu Markdown (.md)',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Beispiele:
  python convert_word_to_markdown.py dokument.docx
  python convert_word_to_markdown.py dokument.docx ausgabe.md
  python convert_word_to_markdown.py --folder ./documents
        """
    )
    
    parser.add_argument(
        'input',
        nargs='?',
        help='Eingabe .docx-Datei'
    )
    
    parser.add_argument(
        'output',
        nargs='?',
        help='Ausgabe .md-Datei (optional, Standard: Eingabedateiname.md)'
    )
    
    parser.add_argument(
        '--folder',
        '-f',
        help='Konvertiert alle .docx-Dateien in einem Ordner'
    )
    
    args = parser.parse_args()
    
    # Batch-Modus
    if args.folder:
        folder_path = Path(args.folder)
        if not folder_path.is_dir():
            print(f"Fehler: {args.folder} ist kein gültiger Ordner")
            return
        
        docx_files = list(folder_path.glob('*.docx'))
        if not docx_files:
            print(f"Keine .docx-Dateien in {args.folder} gefunden")
            return
        
        print(f"Konvertiere {len(docx_files)} Datei(en)...")
        success_count = 0
        for docx_file in docx_files:
            if convert_docx_to_markdown(docx_file):
                success_count += 1
        
        print(f"\nFertig: {success_count}/{len(docx_files)} Dateien erfolgreich konvertiert")
        return
    
    # Einzeldatei-Modus
    if not args.input:
        parser.print_help()
        return
    
    if not os.path.exists(args.input):
        print(f"Fehler: Datei {args.input} nicht gefunden")
        return
    
    convert_docx_to_markdown(args.input, args.output)


if __name__ == '__main__':
    main()

